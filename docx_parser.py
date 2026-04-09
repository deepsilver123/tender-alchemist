# docx_parser.py
"""
Прямой парсинг DOCX с извлечением текста и таблиц.
Без Docling — всё делается локально.
"""

from docx import Document
from bs4 import BeautifulSoup
import re


def extract_from_docx(file_path) -> str:
    """
    Извлекает текст и таблицы из DOCX, формирует простой HTML.
    Таблицы конвертируются в HTML-таблицы без Docling.
    """
    print(f"  Открываю документ...")
    doc = Document(file_path)
    
    print(f"  Параграфов: {len(doc.paragraphs)}, таблиц: {len(doc.tables)}")
    
    result_parts = []
    
    # Обрабатываем всё по порядку
    for element in doc.element.body:
        # Параграф
        if element.tag.endswith('p'):
            for para in doc.paragraphs:
                if para._element == element and para.text.strip():
                    result_parts.append(f"<p>{escape_html(para.text)}</p>")
                    break
        
        # Таблица
        elif element.tag.endswith('tbl'):
            # Находим эту таблицу в doc.tables
            for table in doc.tables:
                if table._element == element:
                    html_table = table_to_html(table)
                    if html_table:
                        result_parts.append(html_table)
                    break
    
    return '\n'.join(result_parts)


def table_to_html(table) -> str:
    """
    Конвертирует python-docx таблицу в HTML.
    Обрабатывает вложенные таблицы рекурсивно.
    """
    rows = len(table.rows)
    cols = len(table.columns)
    
    if rows == 0 or cols == 0:
        return ""
    
    # Начинаем таблицу
    html = ['<table border="1" style="border-collapse: collapse;">']
    
    for i, row in enumerate(table.rows):
        html.append('<tr>')
        
        for j, cell in enumerate(row.cells):
            # Получаем текст ячейки
            cell_text = cell.text.strip()
            
            # Проверяем, есть ли в ячейке вложенная таблица
            nested_tables = cell.tables
            if nested_tables:
                # Если есть вложенная таблица, рекурсивно обрабатываем её
                nested_html = []
                for nested in nested_tables:
                    nested_html.append(table_to_html(nested))
                cell_content = '\n'.join(nested_html)
                # Добавляем также текст ячейки, если есть
                if cell_text:
                    cell_content = f"<p>{escape_html(cell_text)}</p>\n{cell_content}"
            else:
                cell_content = escape_html(cell_text)
            
            html.append(f'<td>{cell_content}</td>')
        
        html.append('</tr>')
    
    html.append('</table>')
    return '\n'.join(html)


def escape_html(text: str) -> str:
    """Экранирует HTML-спецсимволы."""
    if not text:
        return ""
    return (text
            .replace('&', '&amp;')
            .replace('<', '&lt;')
            .replace('>', '&gt;')
            .replace('"', '&quot;')
            .replace("'", '&#39;'))