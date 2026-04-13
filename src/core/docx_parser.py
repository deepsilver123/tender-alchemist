# docx_parser.py (moved into core)
from docx import Document
import logging


def extract_from_docx(file_path) -> str:
    logger = logging.getLogger("tender")
    try:
        logger.info("Открываю документ...")
    except Exception:
        pass
    doc = Document(file_path)
    try:
        logger.info(f"Параграфов: {len(doc.paragraphs)}, таблиц: {len(doc.tables)}")
    except Exception:
        pass
    
    result_parts = []
    
    for element in doc.element.body:
        if element.tag.endswith('p'):
            for para in doc.paragraphs:
                if para._element == element and para.text.strip():
                    result_parts.append(f"<p>{escape_html(para.text)}</p>")
                    break
        elif element.tag.endswith('tbl'):
            for table in doc.tables:
                if table._element == element:
                    html_table = table_to_html(table)
                    if html_table:
                        result_parts.append(html_table)
                    break
    
    return '\n'.join(result_parts)


def table_to_html(table) -> str:
    rows = len(table.rows)
    cols = len(table.columns)
    
    if rows == 0 or cols == 0:
        return ""
    
    html = ['<table border="1" style="border-collapse: collapse;">']
    
    for i, row in enumerate(table.rows):
        html.append('<tr>')
        
        for j, cell in enumerate(row.cells):
            cell_text = cell.text.strip()
            nested_tables = cell.tables
            if nested_tables:
                nested_html = []
                for nested in nested_tables:
                    nested_html.append(table_to_html(nested))
                cell_content = '\n'.join(nested_html)
                if cell_text:
                    cell_content = f"<p>{escape_html(cell_text)}</p>\n{cell_content}"
            else:
                cell_content = escape_html(cell_text)
            
            html.append(f'<td>{cell_content}</td>')
        
        html.append('</tr>')
    
    html.append('</table>')
    return '\n'.join(html)


def escape_html(text: str) -> str:
    if not text:
        return ""
    return (text
            .replace('&', '&amp;')
            .replace('<', '&lt;')
            .replace('>', '&gt;')
            .replace('"', '&quot;')
            .replace("'", '&#39;'))
